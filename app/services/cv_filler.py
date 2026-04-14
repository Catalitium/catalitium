"""Build a Harvard-format CV docx from structured data using the original template styles."""
from __future__ import annotations

import io
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

_TEMPLATE = Path(__file__).parent.parent / "data" / "cv_template.docx"


def _clear_body(doc: Document) -> None:
    """Remove all body content paragraphs, preserving section properties."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _tab_para(doc: Document, left: str, right: str, style: str) -> None:
    """Add paragraph with left text and optional right-aligned content after tab."""
    p = doc.add_paragraph(style=style)
    p.add_run(left)
    if right and right.strip():
        p.add_run("\t" + right)


def render_cv(data: dict) -> bytes:
    """Fill Harvard CV template with structured data, return .docx bytes."""
    doc = Document(_TEMPLATE)
    _clear_body(doc)

    # ── Header ──────────────────────────────────────────────────────────────
    doc.add_paragraph(data.get("name", ""), style="Normal")
    doc.add_paragraph(data.get("contact_line", ""), style="Body Text")

    # ── Education ───────────────────────────────────────────────────────────
    doc.add_paragraph("Education", style="Heading 1")
    for edu in data.get("education", []):
        _tab_para(doc, edu.get("institution", ""), edu.get("location", ""), "Normal")
        p = doc.add_paragraph(style="Body Text")
        p.add_run(edu.get("degree", ""))
        if edu.get("dates", "").strip():
            p.add_run("\t" + edu["dates"])
        if edu.get("coursework", "").strip():
            doc.add_paragraph(
                f"Relevant Coursework: {edu['coursework']}", style="Body Text"
            )
        if edu.get("thesis", "").strip():
            doc.add_paragraph(f"Thesis: {edu['thesis']}", style="Body Text")

    # ── Experience ──────────────────────────────────────────────────────────
    experience = data.get("experience", [])
    if experience:
        doc.add_paragraph("Experience", style="Heading 1")
        for exp in experience:
            _tab_para(doc, exp.get("org", ""), exp.get("location", ""), "Normal")
            _tab_para(doc, exp.get("title", ""), exp.get("dates", ""), "Normal")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Paragraph")

    # ── Leadership & Activities ─────────────────────────────────────────────
    activities = data.get("activities", [])
    if activities:
        doc.add_paragraph("Leadership & Activities", style="Heading 1")
        for act in activities:
            _tab_para(doc, act.get("org", ""), act.get("location", ""), "Normal")
            p = doc.add_paragraph(style="Body Text")
            p.add_run(act.get("role", ""))
            if act.get("dates", "").strip():
                p.add_run("\t" + act["dates"])
            for bullet in act.get("bullets", []):
                doc.add_paragraph(bullet, style="List Paragraph")

    # ── Skills & Interests ──────────────────────────────────────────────────
    skills = data.get("skills", {})
    skill_pairs = [
        ("Technical", skills.get("technical", "")),
        ("Language", skills.get("language", "")),
        ("Laboratory", skills.get("laboratory", "")),
        ("Interests", skills.get("interests", "")),
    ]
    visible = [(k, v) for k, v in skill_pairs if v and v.strip()]
    if visible:
        doc.add_paragraph("Skills & Interests", style="Normal")
        for label, value in visible:
            doc.add_paragraph(f"{label}: {value}", style="Body Text")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
