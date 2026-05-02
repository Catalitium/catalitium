"""CV pipeline: extract text from uploads, map text to structured JSON, render Harvard DOCX."""

from __future__ import annotations

import io
import json
import os
import re
import urllib.request as _req
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader

# =============================================================================
# ── Extract text (PDF / DOCX / pasted text) ──────────────────────────────────
# =============================================================================

ALLOWED_EXTENSIONS = {"pdf", "docx"}
MAX_UPLOAD_BYTES = 4 * 1024 * 1024  # 4MB safeguard for parser memory.
MAX_TEXT_CHARS = 50_000


@dataclass
class ExtractedCV:
    """Normalized extracted CV payload."""

    text: str
    filename: str
    extension: str
    byte_size: int
    original_chars: int
    truncated: bool


class CVExtractionError(ValueError):
    """Typed extraction error with API-friendly metadata."""

    def __init__(self, code: str, message: str, status: int = 400) -> None:
        super().__init__(message)
        self.code = code
        self.message = message
        self.status = status


def extract_cv_from_upload(upload: Any) -> ExtractedCV:
    """Extract and normalize CV text from a Werkzeug FileStorage object."""
    filename = (getattr(upload, "filename", "") or "").strip()
    if not filename:
        raise CVExtractionError("missing_file_name", "Please select a PDF or DOCX CV file.")
    extension = Path(filename).suffix.lower().lstrip(".")
    if extension not in ALLOWED_EXTENSIONS:
        raise CVExtractionError("unsupported_file_type", "Only PDF and DOCX files are supported.")

    try:
        seek = getattr(upload, "seek", None)
        if callable(seek):
            seek(0)
    except Exception:
        pass
    raw = upload.read() or b""
    if not raw:
        raise CVExtractionError("empty_file", "The uploaded file is empty.")
    if len(raw) > MAX_UPLOAD_BYTES:
        raise CVExtractionError("file_too_large", "CV file is too large. Keep it under 4MB.", 413)

    if extension == "pdf":
        parsed = _extract_pdf(raw)
    else:
        parsed = _extract_docx(raw)

    normalized = _normalize_text(parsed)
    if not normalized:
        raise CVExtractionError(
            "no_extractable_text",
            "Could not extract readable text from the file. Try a different CV export.",
        )

    original_chars = len(normalized)
    truncated = original_chars > MAX_TEXT_CHARS
    text = normalized[:MAX_TEXT_CHARS]
    return ExtractedCV(
        text=text,
        filename=filename,
        extension=extension,
        byte_size=len(raw),
        original_chars=original_chars,
        truncated=truncated,
    )


def normalize_cv_text(raw_text: str) -> str:
    """Normalize user-provided text fallback from textarea payloads."""
    text = _normalize_text(raw_text)
    if not text:
        raise CVExtractionError("empty_text", "Please paste CV text or upload a PDF/DOCX file.")
    original_chars = len(text)
    if original_chars > MAX_TEXT_CHARS:
        text = text[:MAX_TEXT_CHARS]
    return text


def _extract_pdf(raw_bytes: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(raw_bytes))
        chunks = []
        for page in reader.pages:
            chunks.append(page.extract_text() or "")
        return "\n".join(chunks)
    except Exception as exc:  # pragma: no cover - parser internals
        raise CVExtractionError("pdf_parse_failed", "Unable to read this PDF CV.") from exc


def _extract_docx(raw_bytes: bytes) -> str:
    try:
        doc = Document(io.BytesIO(raw_bytes))
        paragraphs = [p.text for p in doc.paragraphs]
        return "\n".join(paragraphs)
    except Exception as exc:  # pragma: no cover - parser internals
        raise CVExtractionError("docx_parse_failed", "Unable to read this DOCX CV.") from exc


def _normalize_text(text: str) -> str:
    text = str(text or "").replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# =============================================================================
# ── Structure mapping (Claude / heuristic) ───────────────────────────────────
# =============================================================================

_MODEL = "claude-haiku-4-5-20251001"
_MAX_TOKENS = 2048
_CV_CHARS = 15_000

_SYSTEM = (
    "You are a precise CV parser. Extract structured data and return ONLY valid JSON. "
    "No markdown, no explanation. Rewrite bullet points in Harvard action-verb format: "
    "strong past-tense verb + specific contribution + quantified result where inferable."
)

_SCHEMA = """{
  "name": "Full Name",
  "contact_line": "City, Country · email@example.com · +X XXX-XXXX",
  "education": [
    {
      "institution": "University Name",
      "location": "City, Country",
      "degree": "B.Sc. Computer Science, GPA 3.9",
      "dates": "Month Year",
      "coursework": "Course1, Course2",
      "thesis": "Title"
    }
  ],
  "experience": [
    {
      "org": "Organization Name",
      "location": "City, Country",
      "title": "Job Title",
      "dates": "Month Year – Month Year",
      "bullets": [
        "Action verb + specific contribution + result/metric"
      ]
    }
  ],
  "activities": [
    {
      "org": "Organization or Club Name",
      "location": "City, Country",
      "role": "Role Title",
      "dates": "Month Year – Month Year",
      "bullets": ["..."]
    }
  ],
  "skills": {
    "technical": "Python, SQL, React",
    "language": "English (Native), Spanish (B2)",
    "laboratory": "",
    "interests": "Photography, open-source"
  }
}"""


def extract_cv_structure(cv_text: str) -> dict:
    """Map CV plain text to a Harvard-template-compatible dict.

    Uses Claude Haiku when ANTHROPIC_API_KEY is set, otherwise falls back
    to the heuristic parser.
    """
    api_key = os.getenv("ANTHROPIC_API_KEY", "").strip()
    if api_key:
        return _extract_via_api(cv_text, api_key)
    return _extract_heuristic(cv_text)


def _extract_via_api(cv_text: str, api_key: str) -> dict:
    prompt = (
        f"Parse the following CV into this exact JSON schema:\n{_SCHEMA}\n\n"
        "Rules:\n"
        "- education and experience in reverse chronological order\n"
        "- bullets: rewrite as Harvard action-verb format\n"
        "- contact_line: use · as separator\n"
        "- omit optional fields (coursework, thesis, laboratory) if absent — use empty string\n"
        "- return ONLY JSON, no markdown fences\n\n"
        f"CV TEXT:\n{cv_text[:_CV_CHARS]}"
    )
    payload = {
        "model": _MODEL,
        "max_tokens": _MAX_TOKENS,
        "system": _SYSTEM,
        "messages": [{"role": "user", "content": prompt}],
    }

    text: str | None = None
    try:
        import anthropic  # type: ignore

        client = anthropic.Anthropic(api_key=api_key)
        msg = client.messages.create(**payload)
        text = msg.content[0].text
    except ImportError:
        body = json.dumps(payload).encode()
        req = _req.Request(
            "https://api.anthropic.com/v1/messages",
            data=body,
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
        )
        with _req.urlopen(req, timeout=30) as resp:
            text = json.loads(resp.read())["content"][0]["text"]

    raw = (text or "").strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
        raw = raw.split("```")[0]
    return json.loads(raw.strip())


_SECTION_PATTERNS = {
    "education": re.compile(r"^education", re.I),
    "experience": re.compile(r"^(experience|work experience|employment|professional experience)", re.I),
    "activities": re.compile(r"^(activities|leadership|volunteering|extra.?curricular)", re.I),
    "skills": re.compile(r"^(skills|competencies|technical skills|skills & interests)", re.I),
    "contact": re.compile(r"^(contact|personal details|profile)", re.I),
}
_BULLET_RE = re.compile(r"^[\•\-\*\u2022\u25cf]\s+")
_DATE_RE = re.compile(
    r"(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z]*[\s,]+\d{4}"
    r"|(\d{4})\s*[–\-—]\s*(\d{4}|present|current)",
    re.I,
)


def _strip_bullet(line: str) -> str:
    return _BULLET_RE.sub("", line).strip()


def _is_bullet(line: str) -> bool:
    return bool(_BULLET_RE.match(line))


def _split_sections(text: str) -> dict[str, list[str]]:
    lines = [l.rstrip() for l in text.splitlines()]
    sections: dict[str, list[str]] = {"header": []}
    current = "header"
    for line in lines:
        stripped = line.strip()
        matched = False
        for name, pat in _SECTION_PATTERNS.items():
            if pat.match(stripped) and len(stripped) < 60:
                current = name
                sections.setdefault(current, [])
                matched = True
                break
        if not matched:
            sections.setdefault(current, []).append(line)
    return sections


def _parse_header(lines: list[str]) -> tuple[str, str]:
    non_empty = [l.strip() for l in lines if l.strip()]
    name = non_empty[0] if non_empty else ""
    contact_parts = []
    for l in non_empty[1:]:
        if any(c in l for c in ("@", "+", "·", "|", ",")) or re.search(r"\d{4,}", l):
            contact_parts.append(l)
    contact_line = " · ".join(contact_parts) if contact_parts else " · ".join(non_empty[1:3])
    return name, contact_line


def _parse_entries(lines: list[str]) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        if _is_bullet(stripped):
            if current is None:
                current = {"org": "", "location": "", "title": "", "dates": "", "bullets": []}
                entries.append(current)
            current.setdefault("bullets", []).append(_strip_bullet(stripped))
        else:
            entry_complete = current is not None and bool(current.get("title") or current.get("dates"))
            if current is None or entry_complete:
                current = {"org": stripped, "location": "", "title": "", "dates": "", "bullets": []}
                entries.append(current)
                parts = re.split(r"\s*[,|·]\s*", stripped, maxsplit=1)
                if len(parts) == 2:
                    current["org"] = parts[0].strip()
                    current["location"] = parts[1].strip()
            else:
                date_match = _DATE_RE.search(stripped)
                if date_match:
                    current["dates"] = date_match.group(0)
                    title_part = stripped[: date_match.start()].strip(" –-—")
                    current["title"] = title_part or current.get("title", "")
                else:
                    current["title"] = stripped
    return entries


def _parse_skills(lines: list[str]) -> dict[str, str]:
    skills: dict[str, str] = {"technical": "", "language": "", "laboratory": "", "interests": ""}
    label_map = {
        "technical": "technical",
        "tech": "technical",
        "programming": "technical",
        "software": "technical",
        "language": "language",
        "languages": "language",
        "lab": "laboratory",
        "laboratory": "laboratory",
        "interest": "interests",
        "interests": "interests",
        "hobbies": "interests",
    }
    ungrouped: list[str] = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        colon_pos = stripped.find(":")
        if colon_pos != -1:
            label_raw = stripped[:colon_pos].strip().lower()
            value = stripped[colon_pos + 1 :].strip()
            bucket = label_map.get(label_raw)
            if bucket:
                skills[bucket] = value
                continue
        ungrouped.append(_strip_bullet(stripped) if _is_bullet(stripped) else stripped)
    if ungrouped and not skills["technical"]:
        skills["technical"] = ", ".join(ungrouped)
    return skills


def _extract_heuristic(cv_text: str) -> dict:
    """Best-effort heuristic parser — no network calls."""
    sections = _split_sections(cv_text)
    name, contact_line = _parse_header(sections.get("header", []))

    edu_entries = _parse_entries(sections.get("education", []))
    education = [
        {
            "institution": e.get("org", ""),
            "location": e.get("location", ""),
            "degree": e.get("title", ""),
            "dates": e.get("dates", ""),
            "coursework": "",
            "thesis": "",
        }
        for e in edu_entries
    ]

    exp_entries = _parse_entries(sections.get("experience", []))
    experience = [
        {
            "org": e.get("org", ""),
            "location": e.get("location", ""),
            "title": e.get("title", ""),
            "dates": e.get("dates", ""),
            "bullets": e.get("bullets", []),
        }
        for e in exp_entries
    ]

    act_entries = _parse_entries(sections.get("activities", []))
    activities = [
        {
            "org": e.get("org", ""),
            "location": e.get("location", ""),
            "role": e.get("title", ""),
            "dates": e.get("dates", ""),
            "bullets": e.get("bullets", []),
        }
        for e in act_entries
    ]

    return {
        "name": name,
        "contact_line": contact_line,
        "education": education,
        "experience": experience,
        "activities": activities,
        "skills": _parse_skills(sections.get("skills", [])),
    }


# =============================================================================
# ── Harvard DOCX render ─────────────────────────────────────────────────────
# =============================================================================

_TEMPLATE = Path(__file__).parent / "data" / "cv_template.docx"


def _clear_body(doc: Document) -> None:
    """Remove all body content paragraphs, preserving section properties."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _tab_para(doc: Document, left: str, right: str, style: str) -> None:
    """Add paragraph with left text and optional right-aligned content after tab."""
    p = doc.add_paragraph(style=style)
    p.add_run(left)
    if right and right.strip():
        p.add_run("\t" + right)


def render_cv(data: dict) -> bytes:
    """Fill Harvard CV template with structured data, return .docx bytes."""
    doc = Document(_TEMPLATE)
    _clear_body(doc)

    doc.add_paragraph(data.get("name", ""), style="Normal")
    doc.add_paragraph(data.get("contact_line", ""), style="Body Text")

    doc.add_paragraph("Education", style="Heading 1")
    for edu in data.get("education", []):
        _tab_para(doc, edu.get("institution", ""), edu.get("location", ""), "Normal")
        p = doc.add_paragraph(style="Body Text")
        p.add_run(edu.get("degree", ""))
        if edu.get("dates", "").strip():
            p.add_run("\t" + edu["dates"])
        if edu.get("coursework", "").strip():
            doc.add_paragraph(
                f"Relevant Coursework: {edu['coursework']}", style="Body Text"
            )
        if edu.get("thesis", "").strip():
            doc.add_paragraph(f"Thesis: {edu['thesis']}", style="Body Text")

    experience = data.get("experience", [])
    if experience:
        doc.add_paragraph("Experience", style="Heading 1")
        for exp in experience:
            _tab_para(doc, exp.get("org", ""), exp.get("location", ""), "Normal")
            _tab_para(doc, exp.get("title", ""), exp.get("dates", ""), "Normal")
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Paragraph")

    activities = data.get("activities", [])
    if activities:
        doc.add_paragraph("Leadership & Activities", style="Heading 1")
        for act in activities:
            _tab_para(doc, act.get("org", ""), act.get("location", ""), "Normal")
            p = doc.add_paragraph(style="Body Text")
            p.add_run(act.get("role", ""))
            if act.get("dates", "").strip():
                p.add_run("\t" + act["dates"])
            for bullet in act.get("bullets", []):
                doc.add_paragraph(bullet, style="List Paragraph")

    skills = data.get("skills", {})
    skill_pairs = [
        ("Technical", skills.get("technical", "")),
        ("Language", skills.get("language", "")),
        ("Laboratory", skills.get("laboratory", "")),
        ("Interests", skills.get("interests", "")),
    ]
    visible = [(k, v) for k, v in skill_pairs if v and v.strip()]
    if visible:
        doc.add_paragraph("Skills & Interests", style="Normal")
        for label, value in visible:
            doc.add_paragraph(f"{label}: {value}", style="Body Text")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


__all__ = [
    "ALLOWED_EXTENSIONS",
    "MAX_UPLOAD_BYTES",
    "MAX_TEXT_CHARS",
    "CVExtractionError",
    "ExtractedCV",
    "extract_cv_from_upload",
    "extract_cv_structure",
    "normalize_cv_text",
    "render_cv",
]
